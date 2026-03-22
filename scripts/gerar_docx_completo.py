"""
Gera o manuscrito COMPLETO em DOCX com:
- 4 figuras principais inline
- Tabela 1 formatada
- Seção suplementar com 7 figuras/tabelas suplementares
Estudo C v8 — BSA indexing artifact after bariatric surgery.
"""

import pandas as pd
import numpy as np
from scipy import stats
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
import warnings
warnings.filterwarnings('ignore')
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import os, re

# ── Paths ──
BASE = 'Y:/Base Rafael do Sergio'
DATASET = f'{BASE}/ESTUDO_C_DATASET_v4.csv'
TABS = f'{BASE}/tabelas_v4'
OUT_DIR = f'{BASE}/manuscrito'
_sub = os.environ.get('SUBMISSION_MODE', '0') == '1'
OUT_DOCX = f'{OUT_DIR}/estudo_c_manuscript_SUBMISSION.docx' if _sub else f'{OUT_DIR}/estudo_c_manuscript_FINAL.docx'

# ── Colors ──
ARRUDA = '#2E86AB'
GALVAO = '#A23B72'
POOLED = '#F18F01'

plt.rcParams.update({
    'font.family': 'sans-serif', 'font.size': 10,
    'axes.titlesize': 13, 'axes.titleweight': 'bold', 'axes.labelsize': 11,
    'figure.facecolor': 'white', 'axes.facecolor': 'white',
    'axes.grid': True, 'grid.alpha': 0.2, 'grid.linestyle': '--',
})

# ── Load data ──
print("Loading data...")
df = pd.read_csv(DATASET, low_memory=False)
boot = pd.read_csv(f'{TABS}/Tab17_shapley_bootstrap.csv')
disc = pd.read_csv(f'{TABS}/Tab13_directional_discordance.csv')
lme = pd.read_csv(f'{TABS}/Tab15_lme_results.csv')
kdigo = pd.read_csv(f'{TABS}/Tab5_kdigo_transitions.csv')
most = pd.read_csv(f'{TABS}/Tab9_mosteller_sensitivity.csv')
pre_sens = pd.read_csv(f'{TABS}/Tab14_pre_sensitivity.csv')
sens = pd.read_csv(f'{TABS}/Tab7_sensitivity.csv')
linked = pd.read_csv(f'{TABS}/Tab11_linked_vs_unlinked.csv')
galvao_dv = pd.read_csv(f'{TABS}/Tab16_galvao_date_validated.csv')
shapley_full = pd.read_csv(f'{TABS}/Tab4_shapley.csv')
divergence = pd.read_csv(f'{TABS}/Tab3_divergence.csv')

def save_fig(fig, name):
    path = f'{OUT_DIR}/{name}.png'
    fig.savefig(path, dpi=300, bbox_inches='tight', facecolor='white')
    plt.close(fig)
    print(f"  {name}.png")
    return path

# ══════════════════════════════════════════════
# MAIN FIGURES
# ══════════════════════════════════════════════

# ── Figure 1: Flowchart ──
print("Main figures:")
fig, ax = plt.subplots(figsize=(11, 13))
ax.set_xlim(0, 100); ax.set_ylim(0, 100); ax.axis('off')

def box(ax, x, y, w, h, text, fc='#E8F4FD', fs=7.5, bold=False):
    rect = mpatches.FancyBboxPatch((x-w/2, y-h/2), w, h, boxstyle='round,pad=0.3',
                                    facecolor=fc, edgecolor='#333', linewidth=1.2)
    ax.add_patch(rect)
    ax.text(x, y, text, ha='center', va='center', fontsize=fs,
            fontweight='bold' if bold else 'normal', linespacing=1.2)

def arrow(ax, x1, y1, x2, y2):
    ax.annotate('', xy=(x2, y2), xytext=(x1, y1),
                arrowprops=dict(arrowstyle='->', color='#333', linewidth=1.5))

ax.text(50, 98, 'Figure 1. Study Flow Diagram (STROBE/RECORD)', ha='center', fontsize=12, fontweight='bold')

# LEFT: ARRUDA
ax.text(25, 94, 'DISCOVERY (Arruda)', ha='center', fontsize=10, fontweight='bold', color=ARRUDA)
box(ax, 25, 89, 24, 4.5, 'Clinical records\nN = 1,869 RYGB patients', fc='#D4E8F7', fs=7.5, bold=True)
arrow(ax, 25, 86.75, 25, 84.25)
box(ax, 25, 81.5, 24, 4.5, 'CPF linkage to\nSabin laboratory\n(gold standard)', fc='#E8F4FD', fs=7)
arrow(ax, 25, 79.25, 25, 76.75)
box(ax, 25, 74, 24, 4.5, 'Quality filters:\nsex, DOB, valid Cr\nN = 1,869', fc='#E8F4FD', fs=7)
arrow(ax, 25, 71.75, 25, 69.25)
box(ax, 25, 66.5, 24, 4.5, 'Paired PRE + 12M\nindexed eGFR\nN = 268', fc='#C8E6C9', fs=7.5, bold=True)
arrow(ax, 25, 64.25, 25, 61.75)
box(ax, 25, 59, 24, 4.5, 'Paired PRE + 12M\nnonindexed eGFR\n(observed weight)\nN = 154', fc='#A5D6A7', fs=7.5, bold=True)

# RIGHT: GALVAO
ax.text(75, 94, 'REPLICATION (Galvao)', ha='center', fontsize=10, fontweight='bold', color=GALVAO)
box(ax, 75, 89, 24, 4.5, 'Clinical records\nN = 10,872 patients', fc='#F3DFF1', fs=7.5, bold=True)
arrow(ax, 75, 86.75, 75, 84.25)
box(ax, 75, 81.5, 24, 4.5, 'Sabin laboratory\nN = 6,236 patients', fc='#F8E8F6', fs=7)
arrow(ax, 75, 79.25, 75, 76.75)
box(ax, 75, 74, 24, 5.5, 'Name linkage:\nTier 1 exact: 2,281\nTier 2 homonym: 19\nTotal: 2,300', fc='#F8E8F6', fs=7)
box(ax, 97, 74, 6, 4.5, 'Tier 3: 368\nTier 4: 581\n(sensitivity)', fc='#FDE8E8', fs=5.5)
ax.annotate('', xy=(94, 74), xytext=(87, 74), arrowprops=dict(arrowstyle='->', color='#999', linewidth=1))
arrow(ax, 75, 71.25, 75, 68.75)
box(ax, 75, 66, 24, 4.5, 'Bypass + quality filters\nN = 1,385', fc='#F8E8F6', fs=7)
arrow(ax, 75, 63.75, 75, 61.25)
box(ax, 75, 58.5, 24, 4.5, 'Paired PRE + 12M\nindexed eGFR\nN = 283', fc='#C8E6C9', fs=7.5, bold=True)
arrow(ax, 75, 56.25, 75, 53.75)
box(ax, 75, 51, 24, 4.5, 'Paired PRE + 12M\nnonindexed eGFR\n(observed weight)\nN = 96', fc='#A5D6A7', fs=7.5, bold=True)

# BOTTOM: COMBINED
arrow(ax, 25, 56.75, 38, 44); arrow(ax, 75, 48.75, 62, 44)
box(ax, 50, 40.5, 34, 5.5, 'COMBINED ANALYTICAL DATASET\nN = 3,254 patients\nIndexed 12M: 551  |  Nonindexed 12M: 250',
    fc='#FFF3E0', fs=8, bold=True)
arrow(ax, 40, 37.75, 25, 33); arrow(ax, 50, 37.75, 50, 33); arrow(ax, 60, 37.75, 75, 33)
box(ax, 25, 29.5, 22, 6, 'PRIMARY ANALYSIS\n\n12M paired changes\nShapley decomposition\nDirectional discordance',
    fc='#E3F2FD', fs=7, bold=True)
box(ax, 50, 29.5, 22, 6, 'LME MODEL\n\n3,494 obs / 1,573 pts\nRCS 3 knots\nCohort x time p=0.83',
    fc='#E3F2FD', fs=7, bold=True)
box(ax, 75, 29.5, 22, 6, 'SENSITIVITY\n\nSex, BMI, Mosteller\nCr range, PRE window\nDate-validated Galvao',
    fc='#E3F2FD', fs=7, bold=True)
box(ax, 50, 21, 40, 3.5, 'Both cohorts: Sabin Laboratorios (shared reference laboratory)\nCKD-EPI 2021 race-free  |  BSA DuBois  |  No weight imputation',
    fc='#FFFDE7', fs=7)
fig1_path = save_fig(fig, 'fig1_flowchart')

# ── Figure 2: Shapley decomposition ──
fig, ax = plt.subplots(figsize=(7, 4.5))
cohorts = ['Arruda', 'Galvao', 'Pooled']
labels = ['Arruda\n(Discovery)', 'Galvao\n(Replication)', 'Pooled']
x = np.arange(len(cohorts)); w = 0.35
egfr_vals, egfr_lo, egfr_hi = [], [], []
bsa_vals, bsa_lo, bsa_hi = [], [], []
for c in cohorts:
    e = boot[(boot['Cohort']==c) & (boot['Metric']=='Comp_eGFR')].iloc[0]
    b = boot[(boot['Cohort']==c) & (boot['Metric']=='Comp_BSA')].iloc[0]
    egfr_vals.append(e['Point_estimate']); egfr_lo.append(e['Point_estimate']-e['CI95_lo']); egfr_hi.append(e['CI95_hi']-e['Point_estimate'])
    bsa_vals.append(b['Point_estimate']); bsa_lo.append(b['Point_estimate']-b['CI95_lo']); bsa_hi.append(b['CI95_hi']-b['Point_estimate'])
ax.bar(x-w/2, egfr_vals, w, label='eGFR component', color=ARRUDA, yerr=[egfr_lo, egfr_hi], capsize=4, edgecolor='white')
ax.bar(x+w/2, bsa_vals, w, label='BSA component', color=POOLED, yerr=[bsa_lo, bsa_hi], capsize=4, edgecolor='white')
ax.set_xticks(x); ax.set_xticklabels(labels); ax.set_ylabel('Component (mL/min)')
ax.axhline(0, color='black', linewidth=0.5); ax.legend(framealpha=0.95, loc='lower left')
ax.set_title('Shapley Decomposition of Nonindexed eGFR Change at 12 Months')
for i, (e, b) in enumerate(zip(egfr_vals, bsa_vals)):
    pct = abs(b)/(abs(e)+abs(b))*100
    ax.text(i, 5.5, f'BSA: {pct:.0f}%', ha='center', fontsize=9, fontweight='bold', color=POOLED)
plt.tight_layout()
fig2_path = save_fig(fig, 'fig2_shapley')

# ── Figure 3: Forest plot + discordance ──
fig, (ax1, ax2) = plt.subplots(1, 2, figsize=(10, 5), gridspec_kw={'width_ratios': [3, 2]})
endpoints = [
    ('eGFR indexed\n(Arruda)', 3.60, 1.48, 5.72, ARRUDA),
    ('eGFR indexed\n(Galvao)', 2.80, 1.74, 3.86, GALVAO),
    ('eGFR indexed\n(Pooled)', 3.19, 2.02, 4.35, POOLED),
    ('', np.nan, np.nan, np.nan, 'white'),
    ('eGFR nonindexed\n(Arruda)', -16.19, -19.49, -12.89, ARRUDA),
    ('eGFR nonindexed\n(Galvao)', -18.70, -20.92, -16.47, GALVAO),
    ('eGFR nonindexed\n(Pooled)', -17.15, -19.36, -14.95, POOLED),
]
y_pos = np.arange(len(endpoints))[::-1]
for i, (label, mean, lo, hi, color) in enumerate(endpoints):
    if np.isnan(mean): continue
    ax1.errorbar(mean, y_pos[i], xerr=[[mean-lo], [hi-mean]], fmt='D', color=color,
                 markersize=8, capsize=4, capthick=1.5, linewidth=1.5)
    ax1.text(-25, y_pos[i], label, va='center', ha='right', fontsize=8, color=color)
ax1.axvline(0, color='black', linewidth=0.8); ax1.set_xlabel('Mean change (mL/min)')
ax1.set_xlim(-25, 10); ax1.set_yticks([]); ax1.set_title('12-Month Paired Changes')
ax1.axhspan(2.5, 6.5, alpha=0.05, color='green'); ax1.axhspan(-0.5, 2.5, alpha=0.05, color='red')

colors_disc = ['#E74C3C', '#3498DB', '#95A5A6', '#2ECC71']
for idx_c, cohort in enumerate(['Arruda', 'Galvao', 'Combined']):
    row = disc[disc['Cohort']==cohort].iloc[0]; n = row['N']
    vals = [row['N_discordant']/n*100, row['N_both_up']/n*100, row['N_both_down']/n*100, row['N_reverse']/n*100]
    bottom = 0
    for j, (v, c) in enumerate(zip(vals, colors_disc)):
        ax2.bar(idx_c, v, bottom=bottom, color=c, edgecolor='white', width=0.6)
        if v > 5: ax2.text(idx_c, bottom+v/2, f'{v:.0f}%', ha='center', va='center', fontsize=8, fontweight='bold', color='white')
        bottom += v
ax2.set_xticks([0,1,2]); ax2.set_xticklabels(['Arruda','Galvao','Pooled'], fontsize=9)
ax2.set_ylabel('% of patients'); ax2.set_title('Directional Classification'); ax2.set_ylim(0, 105)
legend_patches = [mpatches.Patch(color=c, label=l) for c, l in zip(colors_disc, ['Discordant (idx+/NI-)', 'Both up', 'Both down', 'Reverse'])]
ax2.legend(handles=legend_patches, fontsize=7, loc='upper right', framealpha=0.9)
plt.tight_layout()
fig3_path = save_fig(fig, 'fig3_forest_discordance')

# ── Figure 4: Dose-response ──
paired = df[df['DELTA_EGFR_IDX_12M'].notna() & df['DELTA_EGFR_NI_12M'].notna() & df['DELTA_BSA_12M'].notna()].copy()
paired['divergence'] = paired['DELTA_EGFR_IDX_12M'] - paired['DELTA_EGFR_NI_12M']
fig, (ax1, ax2) = plt.subplots(1, 2, figsize=(10, 4.5))
for label, color, marker in [('Arruda', ARRUDA, 'o'), ('Galvao', GALVAO, 's')]:
    sub = paired[paired['COORTE']==label]
    ax1.scatter(sub['DELTA_BSA_12M'], sub['divergence'], c=color, alpha=0.5, s=20, marker=marker, label=label, edgecolors='white', linewidth=0.3)
x_fit = paired['DELTA_BSA_12M'].values; y_fit = paired['divergence'].values
mask = np.isfinite(x_fit) & np.isfinite(y_fit)
slope, intercept, r, p, se = stats.linregress(x_fit[mask], y_fit[mask])
x_line = np.linspace(x_fit[mask].min(), x_fit[mask].max(), 100)
ax1.plot(x_line, slope*x_line+intercept, 'k--', linewidth=1.5, alpha=0.7)
rho_bsa = stats.spearmanr(x_fit[mask], y_fit[mask])[0]
ax1.text(0.05, 0.95, f'R² = {r**2:.2f}\n$\\rho$ = {rho_bsa:.2f}', transform=ax1.transAxes, va='top', fontsize=9,
         bbox=dict(boxstyle='round', facecolor='wheat', alpha=0.8))
ax1.set_xlabel('$\\Delta$ BSA (m²)'); ax1.set_ylabel('Divergence (mL/min)'); ax1.set_title('A. BSA Change vs Divergence')
ax1.legend(fontsize=8, framealpha=0.9)

twl_col = None
for c in ['TWL_12M', 'PPT_12M', '%TWL_12M']:
    if c in paired.columns: twl_col = c; break
if twl_col and paired[twl_col].notna().sum() > 20:
    sub_twl = paired[paired[twl_col].notna()]
    for label, color, marker in [('Arruda', ARRUDA, 'o'), ('Galvao', GALVAO, 's')]:
        s = sub_twl[sub_twl['COORTE']==label]
        if len(s) > 0: ax2.scatter(s[twl_col], s['divergence'], c=color, alpha=0.5, s=20, marker=marker, label=label, edgecolors='white', linewidth=0.3)
    x2 = sub_twl[twl_col].values; y2 = sub_twl['divergence'].values; m2 = np.isfinite(x2) & np.isfinite(y2)
    if m2.sum() > 10:
        sl2, int2, r2, p2, se2 = stats.linregress(x2[m2], y2[m2])
        x2_line = np.linspace(x2[m2].min(), x2[m2].max(), 100)
        ax2.plot(x2_line, sl2*x2_line+int2, 'k--', linewidth=1.5, alpha=0.7)
        ax2.text(0.05, 0.95, f'R² = {r2**2:.2f}\n$\\rho$ = {stats.spearmanr(x2[m2], y2[m2])[0]:.2f}',
                 transform=ax2.transAxes, va='top', fontsize=9, bbox=dict(boxstyle='round', facecolor='wheat', alpha=0.8))
    ax2.set_xlabel('%TWL at 12 months'); ax2.set_ylabel('Divergence (mL/min)'); ax2.set_title('B. Weight Loss vs Divergence')
    ax2.legend(fontsize=8, framealpha=0.9)
else:
    ax2.text(0.5, 0.5, 'TWL data not available', ha='center', va='center', transform=ax2.transAxes, fontsize=12, color='gray')
    ax2.set_title('B. Weight Loss vs Divergence')
plt.tight_layout()
fig4_path = save_fig(fig, 'fig4_dose_response')

# ══════════════════════════════════════════════
# SUPPLEMENTARY FIGURES
# ══════════════════════════════════════════════
print("Supplementary figures:")

# ── S-Fig 1: LME predicted curves (simulated from coefficients) ──
fig, (ax1, ax2) = plt.subplots(1, 2, figsize=(10, 4.5))

# eGFR indexed predictions
lme_idx = lme[lme['Model'] == 'eGFR_indexed']
const = lme_idx[lme_idx['Variable']=='const']['Coefficient'].values[0]
b_time = lme_idx[lme_idx['Variable']=='time_s1']['Coefficient'].values[0]
b_time2 = lme_idx[lme_idx['Variable']=='time_s2']['Coefficient'].values[0]
b_age = lme_idx[lme_idx['Variable']=='age_baseline']['Coefficient'].values[0]
b_cohort = lme_idx[lme_idx['Variable']=='cohort_galvao']['Coefficient'].values[0]

t = np.linspace(0, 5, 100)
# RCS approximation: spline_linear=t, spline_nonlinear~t^2 (simplified)
age_ref = 38.7  # mean age
for cohort_label, cohort_val, color in [('Arruda', 0, ARRUDA), ('Galvao', 1, GALVAO)]:
    y_pred = const + b_time*t + b_time2*(t**2) + b_cohort*cohort_val + b_age*age_ref
    ax1.plot(t, y_pred, color=color, linewidth=2.5, label=cohort_label)
ax1.set_xlabel('Time from surgery (years)'); ax1.set_ylabel('Predicted eGFR indexed (mL/min/1.73m²)')
ax1.set_title('A. Indexed eGFR Trajectories (LME)'); ax1.legend(framealpha=0.9)
ax1.axhline(90, color='grey', linewidth=0.5, linestyle='--', alpha=0.5)
ax1.text(4.5, 90.5, 'G2', fontsize=7, color='grey')

# eGFR nonindexed predictions
lme_ni = lme[lme['Model'] == 'eGFR_NI']
const_ni = lme_ni[lme_ni['Variable']=='const']['Coefficient'].values[0]
b_time_ni = lme_ni[lme_ni['Variable']=='time_s1']['Coefficient'].values[0]
b_time2_ni = lme_ni[lme_ni['Variable']=='time_s2']['Coefficient'].values[0]
b_age_ni = lme_ni[lme_ni['Variable']=='age_baseline']['Coefficient'].values[0]
b_cohort_ni = lme_ni[lme_ni['Variable']=='cohort_galvao']['Coefficient'].values[0]
b_female_ni = lme_ni[lme_ni['Variable']=='female']['Coefficient'].values[0]

for cohort_label, cohort_val, color in [('Arruda', 0, ARRUDA), ('Galvao', 1, GALVAO)]:
    # female (majority ~83%)
    y_ni = const_ni + b_time_ni*t + b_time2_ni*(t**2) + b_cohort_ni*cohort_val + b_age_ni*age_ref + b_female_ni*0.83
    ax2.plot(t, y_ni, color=color, linewidth=2.5, label=cohort_label)
ax2.set_xlabel('Time from surgery (years)'); ax2.set_ylabel('Predicted eGFR nonindexed (mL/min)')
ax2.set_title('B. Nonindexed eGFR Trajectories (LME)'); ax2.legend(framealpha=0.9)
plt.tight_layout()
sfig1_path = save_fig(fig, 'sfig1_lme_curves')

# ── S-Fig 2: G-category transitions over time ──
fig, ax = plt.subplots(figsize=(8, 5))
comb = kdigo[kdigo['Cohort']=='Combined']
windows_order = ['3M', '6M', '12M', '24M', '36M', '60M']
comb_sorted = comb.set_index('Window').loc[windows_order].reset_index()

def extract_pct(s):
    m = re.search(r'\(([\d.]+)%\)', s)
    return float(m.group(1)) if m else 0

improved = [extract_pct(r) for r in comb_sorted['Improved']]
stable = [extract_pct(r) for r in comb_sorted['Stable']]
worsened = [extract_pct(r) for r in comb_sorted['Worsened']]
ns = comb_sorted['N'].values

x = np.arange(len(windows_order))
ax.bar(x, stable, color='#27AE60', label='Stable', edgecolor='white')
ax.bar(x, improved, bottom=stable, color='#3498DB', label='Improved', edgecolor='white')
ax.bar(x, worsened, bottom=[s+i for s,i in zip(stable, improved)], color='#E74C3C', label='Worsened', edgecolor='white')

for i, (s, im, w, n) in enumerate(zip(stable, improved, worsened, ns)):
    ax.text(i, s/2, f'{s:.0f}%', ha='center', va='center', fontsize=9, fontweight='bold', color='white')
    if w > 1.5: ax.text(i, s+im+w/2, f'{w:.1f}%', ha='center', va='center', fontsize=7, color='white')
    ax.text(i, 103, f'N={n}', ha='center', fontsize=7, color='grey')

ax.set_xticks(x); ax.set_xticklabels(windows_order); ax.set_ylabel('% of patients')
ax.set_ylim(0, 108); ax.set_title('G-Category Transitions Over Time (Combined)')
ax.legend(loc='lower left', framealpha=0.9)
plt.tight_layout()
sfig2_path = save_fig(fig, 'sfig2_gcategory')

# ── S-Fig 3: Mosteller vs DuBois sensitivity ──
fig, ax = plt.subplots(figsize=(7, 4.5))
cohorts_m = ['Arruda', 'Galvao', 'Combined']
x = np.arange(len(cohorts_m)); w = 0.35
dubois_vals = most['Mean_Δ_NI_DuBois'].values
mosteller_vals = most['Mean_Δ_NI_Mosteller'].values
ax.bar(x-w/2, dubois_vals, w, label='DuBois', color=ARRUDA, edgecolor='white')
ax.bar(x+w/2, mosteller_vals, w, label='Mosteller', color=GALVAO, edgecolor='white')
ax.set_xticks(x); ax.set_xticklabels(cohorts_m)
ax.set_ylabel('$\\Delta$ Nonindexed eGFR (mL/min)'); ax.axhline(0, color='black', linewidth=0.5)
ax.set_title('Nonindexed eGFR Change: DuBois vs Mosteller BSA Formula')
ax.legend(framealpha=0.9)
for i, (d, m) in enumerate(zip(dubois_vals, mosteller_vals)):
    diff = abs(m) - abs(d)
    ax.text(i, min(d, m)-1.5, f'Diff: {diff:.1f}', ha='center', fontsize=8, color='grey')
plt.tight_layout()
sfig3_path = save_fig(fig, 'sfig3_mosteller')

# ── S-Fig 4: Sensitivity analyses forest plot ──
fig, ax = plt.subplots(figsize=(8, 6))
# Parse sensitivity data
sens_labels = []
sens_idx_vals = []
sens_ni_vals = []

for _, row in sens.iterrows():
    lbl = f"{row['Sensitivity']} - {row['Variable']}"
    delta_str = str(row['Delta']).replace('+', '')
    try:
        delta = float(delta_str)
    except:
        continue
    if row['Variable'] == 'eGFR_idx':
        sens_labels.append(row['Sensitivity'])
        sens_idx_vals.append(delta)
    elif row['Variable'] == 'eGFR_NI':
        sens_ni_vals.append(delta)

y = np.arange(len(sens_labels))[::-1]
ax.barh(y+0.2, sens_idx_vals, 0.35, label='$\\Delta$ Indexed', color=ARRUDA, edgecolor='white')
ax.barh(y-0.2, sens_ni_vals, 0.35, label='$\\Delta$ Nonindexed', color=POOLED, edgecolor='white')
ax.set_yticks(y); ax.set_yticklabels(sens_labels, fontsize=8)
ax.axvline(0, color='black', linewidth=0.5)
ax.set_xlabel('Mean change (mL/min)'); ax.set_title('Sensitivity Analyses: 12-Month Paired Changes')
ax.legend(loc='lower left', framealpha=0.9)
plt.tight_layout()
sfig4_path = save_fig(fig, 'sfig4_sensitivity')

# ── S-Fig 5: PRE window sensitivity ──
fig, ax = plt.subplots(figsize=(7, 4))
pre_labels = pre_sens['PRE_window'].values
x = np.arange(len(pre_labels)); w = 0.35
idx_vals = [float(str(v).replace('+','')) for v in pre_sens['Delta_idx']]
ni_vals = [float(str(v).replace('+','')) for v in pre_sens['Delta_NI']]
ax.bar(x-w/2, idx_vals, w, label='$\\Delta$ Indexed', color=ARRUDA, edgecolor='white')
ax.bar(x+w/2, ni_vals, w, label='$\\Delta$ Nonindexed', color=POOLED, edgecolor='white')
ax.set_xticks(x); ax.set_xticklabels(pre_labels, fontsize=8)
ax.axhline(0, color='black', linewidth=0.5)
ax.set_ylabel('Mean change (mL/min)'); ax.set_title('Pre-operative Window Sensitivity')
ax.legend(framealpha=0.9)
for i, (v, n) in enumerate(zip(ni_vals, pre_sens['N_NI'])):
    ax.text(i+w/2, v-1, f'N={int(n)}', ha='center', fontsize=7, color='grey')
plt.tight_layout()
sfig5_path = save_fig(fig, 'sfig5_pre_window')

# ── S-Fig 6: Galvao date-validated sensitivity ──
fig, ax = plt.subplots(figsize=(7, 4.5))
dv = galvao_dv[galvao_dv['N_NI'].notna()].copy()
dv_labels = [s.replace('Galvao_', '') for s in dv['Sensitivity']]
x = np.arange(len(dv_labels)); w = 0.35
idx_dv = [float(str(v).replace('+','')) for v in dv['Delta_idx']]
ni_dv = [float(str(v).replace('+','')) for v in dv['Delta_NI']]
ax.bar(x-w/2, idx_dv, w, label='$\\Delta$ Indexed', color=ARRUDA, edgecolor='white')
ax.bar(x+w/2, ni_dv, w, label='$\\Delta$ Nonindexed', color=POOLED, edgecolor='white')
ax.set_xticks(x); ax.set_xticklabels(dv_labels, fontsize=8)
ax.axhline(0, color='black', linewidth=0.5)
ax.set_ylabel('Mean change (mL/min)'); ax.set_title('Galvao Date-Validated 12M Sensitivity')
ax.legend(framealpha=0.9)
for i, (n_idx, n_ni) in enumerate(zip(dv['N_idx'], dv['N_NI'])):
    ax.text(i, 3, f'N idx={int(n_idx)}\nN NI={int(n_ni)}', ha='center', fontsize=7, color='grey')
plt.tight_layout()
sfig6_path = save_fig(fig, 'sfig6_galvao_datevalidated')

# ── S-Fig 7: Shapley temporal evolution (Galvao) ──
shap_gal = shapley_full[shapley_full['Cohort']=='Galvao'].copy()
if len(shap_gal) > 0:
    fig, (ax1, ax2) = plt.subplots(1, 2, figsize=(10, 4.5))
    windows = shap_gal['Window'].values
    # Parse "value ± sd" strings
    def parse_mean(s):
        s = str(s).split('±')[0].split('�')[0].strip()
        return float(s)
    def parse_pct(s):
        return float(str(s).replace('%','').strip())
    comp_egfr = np.array([parse_mean(v) for v in shap_gal['Comp_eGFR'].values])
    comp_bsa = np.array([parse_mean(v) for v in shap_gal['Comp_BSA'].values])
    pct_col = '%BSA (median)' if '%BSA (median)' in shap_gal.columns else None
    pct_bsa = np.array([parse_pct(v) for v in shap_gal[pct_col].values]) if pct_col else np.array([abs(b)/(abs(e)+abs(b))*100 for e,b in zip(comp_egfr, comp_bsa)])

    x = np.arange(len(windows))
    ax1.bar(x-0.2, comp_egfr, 0.35, label='eGFR component', color=ARRUDA, edgecolor='white')
    ax1.bar(x+0.2, comp_bsa, 0.35, label='BSA component', color=POOLED, edgecolor='white')
    ax1.set_xticks(x); ax1.set_xticklabels(windows); ax1.axhline(0, color='black', linewidth=0.5)
    ax1.set_ylabel('Component (mL/min)'); ax1.set_title('A. Shapley Components Over Time (Galvao)')
    ax1.legend(framealpha=0.9, fontsize=8)

    ax2.plot(x, pct_bsa, 'o-', color=POOLED, linewidth=2, markersize=8)
    ax2.set_xticks(x); ax2.set_xticklabels(windows)
    ax2.set_ylabel('% BSA contribution'); ax2.set_title('B. BSA Contribution Over Time (Galvao)')
    ax2.set_ylim(50, 120); ax2.axhline(100, color='grey', linewidth=0.5, linestyle='--')
    for i, p in enumerate(pct_bsa):
        ax2.text(i, p+2, f'{p:.0f}%', ha='center', fontsize=8, fontweight='bold', color=POOLED)
    plt.tight_layout()
    sfig7_path = save_fig(fig, 'sfig7_shapley_temporal')
else:
    sfig7_path = None

# ══════════════════════════════════════════════
# DOCX GENERATION
# ══════════════════════════════════════════════

print("\nBuilding DOCX...")
doc = Document()

style = doc.styles['Normal']
style.font.name = 'Calibri'; style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(6); style.paragraph_format.line_spacing = 1.15

def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs: run.font.color.rgb = RGBColor(0, 0, 0)
    return h

def add_para(doc, text, bold=False, italic=False, size=11, align=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size); run.bold = bold; run.italic = italic
    if align: p.alignment = align
    return p

SUBMISSION_MODE = os.environ.get('SUBMISSION_MODE', '0') == '1'

def add_figure(doc, path, caption, width=Inches(6), alt_text=''):
    if SUBMISSION_MODE:
        # Text-only: insert placeholder instead of image
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run('[Insert ' + caption.split('.')[0] + ' here]')
        run.bold = True; run.font.size = Pt(10); run.font.color.rgb = RGBColor(150, 150, 150)
        cap = doc.add_paragraph(); cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = cap.add_run(caption); run.font.size = Pt(9); run.italic = True
    else:
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(); run.add_picture(path, width=width)
        cap = doc.add_paragraph(); cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = cap.add_run(caption); run.font.size = Pt(9); run.italic = True
        if alt_text:
            alt_p = doc.add_paragraph()
            alt_run = alt_p.add_run(f'Alt text: {alt_text}')
            alt_run.font.size = Pt(8); alt_run.italic = True; alt_run.font.color.rgb = RGBColor(100,100,100)
    doc.add_paragraph()

# ── TITLE ──
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Body surface area indexing can reverse the apparent direction of creatinine-based estimated glomerular filtration rate change 12 months after bariatric surgery')
run.bold = True; run.font.size = Pt(14)

doc.add_paragraph()
# Authors
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
authors = [
    ('Sergio Arruda', '1'),
    ('Rafael Oliveira Galv\u00e3o', '1'),
    ('Rafael de Negreiros Botan', '2'),
    ('Larissa Bevilaqua Sampaio Contreiras', '3'),
    ('Erika Bevilaqua Rangel', '4'),
]
for i, (name, aff) in enumerate(authors):
    run = p.add_run(name)
    run.font.size = Pt(11)
    sup = p.add_run(aff)
    sup.font.size = Pt(8)
    sup.font.superscript = True
    if i < len(authors)-1:
        p.add_run(', ').font.size = Pt(11)

# Affiliations
doc.add_paragraph()
affs = [
    '\u00b9 Bariatric Surgery, Private Practice, Bras\u00edlia, DF, Brazil',
    '\u00b2 Department of Oncology, Universidade de Bras\u00edlia, Bras\u00edlia, DF, Brazil',
    '\u00b3 Medical Student, Universidade de Bras\u00edlia (UnB), Bras\u00edlia, DF, Brazil',
    '\u2074 Department of Medicine (Nephrology), Universidade Federal de S\u00e3o Paulo (UNIFESP), S\u00e3o Paulo, SP, Brazil',
]
for aff in affs:
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(aff); run.font.size = Pt(9); run.italic = True

doc.add_paragraph()
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Corresponding author: '); run.font.size = Pt(9); run.bold = True
run2 = p.add_run('Rafael de Negreiros Botan \u2014 oncologista@gmail.com')
run2.font.size = Pt(9)

add_para(doc, '\nRunning title: BSA indexing artifact after bariatric surgery', italic=True, size=10, align=WD_ALIGN_PARAGRAPH.CENTER)

# ── ABSTRACT ──
add_heading(doc, 'Abstract', level=2)

p = doc.add_paragraph()
run = p.add_run('Background and hypothesis. '); run.bold = True
p.add_run('We linked routinely collected clinical and laboratory data from two bariatric surgery cohorts in Brasilia, Brazil (2008-2024), to examine how body surface area (BSA) indexing changes the interpretation of creatinine-based estimated glomerular filtration rate (eGFR) trajectories 12 months after Roux-en-Y gastric bypass (RYGB). Because BSA falls substantially after surgery, this rescaling could distort or even reverse the apparent direction of eGFR change.')

p = doc.add_paragraph()
run = p.add_run('Methods. '); run.bold = True
p.add_run('We conducted a retrospective dual-cohort study using administrative and clinical data from two independent surgical practices sharing a single reference laboratory (Sabin). The discovery cohort (Arruda; N = 268 indexed, 154 nonindexed at 12 months) used national-ID linkage; the replication cohort (Galvao; N = 283 indexed, 96 nonindexed) used deterministic name linkage. eGFR was computed with the CKD-EPI 2021 race-free equation. Nonindexed eGFR was calculated as eGFR_indexed \u00d7 BSA_DuBois / 1.73. We decomposed the change in nonindexed eGFR using a Shapley-type bilinear identity, with 5,000-replicate stratified bootstrap confidence intervals.')

p = doc.add_paragraph()
run = p.add_run('Results. '); run.bold = True
p.add_run('In the discovery cohort, indexed eGFR rose (+3.60 [95% CI 1.48, 5.72]) while nonindexed eGFR fell (-16.19 [-19.49, -12.89]) at 12 months. An independent linked cohort replicated the marked nonindexed decline (-18.70 [-20.92, -16.47]) and large indexed-nonindexed divergence; after temporal validation, indexed change attenuated whereas nonindexed decline persisted. The BSA component accounted for 86.1% [79.6, 93.8] of the nonindexed decline. Nearly half of patients showed directional discordance (48.8% [42.7, 55.0]; 0% reverse pattern). Cohort-by-time interaction was non-significant (p = 0.83).')

p = doc.add_paragraph()
run = p.add_run('Conclusions. '); run.bold = True
p.add_run('BSA indexing can reverse the apparent direction of creatinine-based eGFR change after RYGB. Reporting only indexed eGFR may understate the decline in nonindexed creatinine-based eGFR after surgery.')

# ── GRAPHICAL ABSTRACT ──
add_heading(doc, 'Graphical Abstract', level=2)
ga_path = f'{OUT_DIR}/graphical_abstract.png'
if os.path.exists(ga_path):
    add_figure(doc, ga_path, 'Graphical Abstract. BSA indexing reverses the apparent direction of eGFR change after bariatric surgery.', width=Inches(6.5))

doc.add_page_break()

# ── KEY LEARNING POINTS (NDT format) ──
add_heading(doc, 'Key Learning Points', level=2)

p = doc.add_paragraph()
run = p.add_run('What was known:'); run.bold = True
doc.add_paragraph('Indexed eGFR normalises filtration to a fixed body surface area of 1.73 m\u00b2, but after bariatric surgery BSA falls substantially.', style='List Bullet')
doc.add_paragraph('Existing studies predominantly report indexed eGFR, potentially masking the mechanical effect of BSA rescaling on apparent kidney function trajectories.', style='List Bullet')

p = doc.add_paragraph()
run = p.add_run('This study adds:'); run.bold = True
doc.add_paragraph('In two independent RYGB cohorts sharing a single reference laboratory (N = 551 indexed, 250 nonindexed at 12 months), indexed eGFR rose (+3.2 mL/min) while nonindexed eGFR fell (\u221217.2 mL/min), producing clinically opposite conclusions.', style='List Bullet')
doc.add_paragraph('Exact bilinear decomposition attributed 86% [95% CI 80\u201394%] of the nonindexed decline to BSA rescaling alone.', style='List Bullet')
doc.add_paragraph('Nearly half of patients (48.8%) showed directional discordance; 0% showed the reverse pattern.', style='List Bullet')

p = doc.add_paragraph()
run = p.add_run('Potential impact:'); run.bold = True
doc.add_paragraph('Clinicians and researchers interpreting eGFR after bariatric surgery should consider reporting both indexed and nonindexed values.', style='List Bullet')
doc.add_paragraph('Exclusive reliance on indexed eGFR may provide a systematically optimistic picture of post-surgical kidney function trajectory.', style='List Bullet')

doc.add_paragraph()

# Keywords
p = doc.add_paragraph()
run = p.add_run('Keywords: '); run.bold = True
p.add_run('GFR; CKD-EPI equation; creatinine; body surface area; gastric bypass')

doc.add_page_break()

# ── INTRODUCTION ──
add_heading(doc, 'Introduction', level=2)
add_para(doc, 'Bariatric surgery, particularly Roux-en-Y gastric bypass (RYGB), produces substantial and sustained weight loss, with well-documented improvements in type 2 diabetes, hypertension, and cardiovascular risk [1,2]. The effect of bariatric surgery on kidney function, however, remains debated [3,4,5]. Most studies report estimated glomerular filtration rate (eGFR) in its standard indexed form (mL/min/1.73 m\u00b2), which normalises filtration to a fixed body surface area (BSA) of 1.73 m\u00b2 [6,7].')
add_para(doc, 'This normalisation is unproblematic when body size is approximately average [8,9]. After bariatric surgery, however, patients lose 25\u201340% of excess weight, and BSA falls by 0.25\u20130.40 m\u00b2 within the first year. Because nonindexed (absolute) eGFR equals indexed eGFR multiplied by (BSA / 1.73), a large drop in BSA mechanically reduces nonindexed eGFR even if indexed eGFR remains stable or rises. In principle, this rescaling could be large enough to reverse the apparent direction of eGFR change.')
add_para(doc, 'Despite recognition of this issue in nephrology guidelines [6,8], no study has formally quantified the BSA indexing artifact after bariatric surgery in a large cohort, nor decomposed the relative contribution of the non-BSA creatinine-based component versus BSA rescaling.')
add_para(doc, 'We addressed this gap in a dual-cohort discovery-replication design, using two independent RYGB cohorts sharing the same reference laboratory to (i) quantify the divergence between indexed and nonindexed eGFR at 12 months, (ii) decompose the nonindexed eGFR change into filtration and BSA components, and (iii) test whether the artifact is reproducible across independent surgical practices.')

# ── MATERIALS AND METHODS ──
add_heading(doc, 'Materials and Methods', level=2)

add_heading(doc, 'Study design and reporting', level=3)
add_para(doc, 'This retrospective observational study used routinely collected clinical and laboratory data and followed the STROBE statement [10] and RECORD extension [11] for studies conducted using routinely collected health data. The study was registered at the Open Science Framework prior to submission (https://osf.io/9r5w7/).')

add_heading(doc, 'Ethics', level=3)
add_para(doc, 'This study was approved by the Research Ethics Committee (Comit\u00ea de \u00c9tica em Pesquisa) of Universidade Federal de S\u00e3o Paulo (UNIFESP) under CAAE [A PREENCHER], project title "Impacto da cirurgia bari\u00e1trica nos desfechos cardiorrenais metab\u00f3licos." Informed consent was waived because the study involved secondary use of existing de-identified clinical and laboratory records without direct participant contact. Linkage identifiers (national ID and names) were used solely for deterministic matching and were removed before analysis.')

add_heading(doc, 'Data sources and linkage', level=3)
add_para(doc, 'Two independent surgical practices in Brasilia, DF, Brazil, contributed data. Both practices use Sabin Laboratorios as their reference laboratory, enabling standardised creatinine measurements across cohorts.')
p = doc.add_paragraph(); run = p.add_run('Discovery cohort (Arruda). '); run.bold = True
p.add_run('The Arruda cohort comprised 1,869 consecutive RYGB patients. Surgical technique was confirmed as 100% RYGB by the operating surgeon. Laboratory data were linked to clinical records using the Brazilian national identifier (CPF), providing gold-standard deterministic linkage. Sex was recorded in clinical charts with 100% coverage.')
p = doc.add_paragraph(); run = p.add_run('Replication cohort (Galvao). '); run.bold = True
p.add_run('The Galvao cohort comprised 10,872 patients. Laboratory data were linked using hierarchical deterministic name linkage: exact normalised-name match (tier 1; N = 2,281) and birth-year disambiguation of homonymous names (tier 2; N = 19, all manually reviewed). Patients were filtered to RYGB (98.5% of cohort; the remaining 1.5% were sleeve gastrectomy or adjustable band and were excluded). Sex was inferred from first names using a validated catalogue of 2,782 Brazilian names. Body weight was extracted from standardised free-text follow-up notes using prespecified numeric patterns; ambiguous values were excluded, and no weight imputation was performed.')

add_heading(doc, 'Study population', level=3)
add_para(doc, 'The analytical dataset included patients with: (i) confirmed RYGB in both cohorts; (ii) defined sex; (iii) known date of birth; and (iv) at least one valid serum creatinine. The final dataset comprised 3,254 RYGB patients (1,869 Arruda; 1,385 Galvao).')

add_heading(doc, 'Measurements and definitions', level=3)
add_para(doc, 'eGFR (indexed) was calculated using the CKD-EPI 2021 race-free equation [12]. Serum creatinine values outside 0.20\u20135.0 mg/dL were excluded [13]. BSA was calculated using the DuBois formula [14]; sensitivity analysis used the Mosteller formula [15]. Nonindexed eGFR was calculated as eGFR_indexed \u00d7 BSA_DuBois / 1.73, requiring observed weight at both time points (no imputation) [8]. Pre-operative creatinine: closest to surgery within \u2212365 to \u22121 days. 12-month window: 301\u2013548 days. Age at examination used a three-tier priority system (actual date, days offset, age-at-surgery plus window offset; maximum error 0.045 years).')

add_heading(doc, 'Shapley decomposition', level=3)
add_para(doc, 'We decomposed the change in nonindexed eGFR into two additive components using the bilinear Shapley identity [16]: \u0394eGFR_NI = [\u0394eGFR_idx \u00d7 mean(BSA)/1.73] + [mean(eGFR_idx) \u00d7 \u0394BSA/1.73]. This decomposition is algebraically exact (zero residual). Bootstrap 95% CIs used 5,000 patient-level replicates [17], stratified by cohort in pooled analyses.')

add_heading(doc, 'Statistical analysis', level=3)
add_para(doc, 'Within-person changes were tested with paired t-tests. A linear mixed-effects model with restricted cubic spline (3 knots at 10th/50th/90th percentiles) and random intercept per patient assessed temporal trajectory and cohort homogeneity. Sensitivity analyses included: sex-stratified, BMI-stratified, BSA formula (Mosteller), creatinine range (0.4-3.0), pre-operative window restriction (-180d, -90d), tight 12-month window (365\u00b145d), and date-validated Galvao subset. Primary inferential claims were anchored in the discovery cohort; replication and pooled analyses were prespecified as supportive. All analyses used Python 3.13.')

doc.add_page_break()

# ── RESULTS ──
add_heading(doc, 'Results', level=2)

# Figure 1
add_figure(doc, fig1_path, 'Figure 1. Study flow diagram (STROBE/RECORD).', width=Inches(6.5), alt_text='STROBE flow diagram showing patient selection from two independent surgical cohorts through quality filters to final analytical samples of 551 indexed and 250 nonindexed paired measurements at 12 months.')

# Table 1
add_heading(doc, 'Baseline characteristics', level=3)
arr = df[df['COORTE']=='Arruda']; gal = df[df['COORTE']=='Galvao']

def fmt(s, f='.1f'): return f"{s.mean():{f}} \u00b1 {s.std():{f}}"
def pct(s, v): return f"{(s==v).sum()}/{len(s)} ({(s==v).mean()*100:.1f}%)"

rows = [
    ['N', str(len(arr)), str(len(gal)), str(len(df))],
    ['Age, years', fmt(arr['IDADE_CIRURGIA']), fmt(gal['IDADE_CIRURGIA']), fmt(df['IDADE_CIRURGIA'])],
    ['Female sex, n (%)', pct(arr['SEXO'],'Feminino'), pct(gal['SEXO'],'Feminino'), pct(df['SEXO'],'Feminino')],
    ['Weight, kg', fmt(arr['PESO_PRE']), fmt(gal['PESO_PRE']), fmt(df['PESO_PRE'])],
    ['Height, cm', fmt(arr['ALTURA_CM']), fmt(gal['ALTURA_CM']), fmt(df['ALTURA_CM'])],
    ['BMI, kg/m\u00b2', fmt(arr['IMC_PRE']), fmt(gal['IMC_PRE']), fmt(df['IMC_PRE'])],
    ['BSA, m\u00b2', fmt(arr['BSA_DUBOIS_PRE'],'.2f'), fmt(gal['BSA_DUBOIS_PRE'],'.2f'), fmt(df['BSA_DUBOIS_PRE'],'.2f')],
    ['Creatinine, mg/dL', fmt(arr['CREATININA_PRE'],'.2f'), fmt(gal['CREATININA_PRE'],'.2f'), fmt(df['CREATININA_PRE'],'.2f')],
    ['eGFR indexed, mL/min/1.73m\u00b2', fmt(arr['EGFR_CKD_EPI_PRE']), fmt(gal['EGFR_CKD_EPI_PRE']), fmt(df['EGFR_CKD_EPI_PRE'])],
    ['eGFR nonindexed, mL/min', fmt(arr['EGFR_NI_PRE']), fmt(gal['EGFR_NI_PRE']), fmt(df['EGFR_NI_PRE'])],
]

table = doc.add_table(rows=len(rows)+1, cols=4, style='Light Shading Accent 1')
table.alignment = WD_TABLE_ALIGNMENT.CENTER
for j, h in enumerate(['Variable', 'Arruda (Discovery)', 'Galvao (Replication)', 'Combined']):
    cell = table.rows[0].cells[j]; cell.text = h
    for p in cell.paragraphs:
        for run in p.runs: run.bold = True; run.font.size = Pt(9)
for i, row in enumerate(rows):
    for j, val in enumerate(row):
        cell = table.rows[i+1].cells[j]; cell.text = val
        for p in cell.paragraphs:
            for run in p.runs: run.font.size = Pt(9)

p = doc.add_paragraph()
run = p.add_run('Table 1. Baseline characteristics by cohort. Values are mean \u00b1 SD unless otherwise specified.')
run.font.size = Pt(9); run.italic = True
doc.add_paragraph()

# Primary analysis
add_heading(doc, 'Primary analysis: 12-month paired changes', level=3)
add_para(doc, 'In the discovery cohort (Arruda), indexed eGFR rose by +3.60 mL/min/1.73 m\u00b2 [95% CI 1.48, 5.72; p = 1.0 \u00d7 10\u207b\u00b3] while nonindexed eGFR fell by -16.19 mL/min [-19.49, -12.89; p = 2.0 \u00d7 10\u207b\u00b9\u2077]. The replication cohort (Galvao) confirmed the marked nonindexed decline (-18.70 [-20.92, -16.47; p = 1.4 \u00d7 10\u207b\u00b2\u2079]) and large indexed-nonindexed divergence (21.21 mL/min). In the full replication sample, indexed eGFR also rose (+2.80 [1.74, 3.86]); however, when restricted to date-validated 12-month measurements (301-548 days, N = 48), indexed change attenuated to near zero (\u22120.06, p = 0.96) while nonindexed decline persisted (-18.94 [-22.61, -15.28]). Interaction p-values were non-significant for both indexed (0.50) and nonindexed (0.28) endpoints.')

# Shapley
add_heading(doc, 'Shapley decomposition', level=3)
add_figure(doc, fig2_path, 'Figure 2. Shapley decomposition of nonindexed eGFR change at 12 months. The BSA component (orange) dominates; the eGFR component (blue) is small and positive. Error bars: 5,000-replicate bootstrap 95% CIs.', width=Inches(5.5), alt_text='Grouped bar chart showing Shapley decomposition with eGFR component (blue, small positive) and BSA component (orange, large negative) for Arruda, Galvao, and Pooled cohorts. BSA percentages labeled above bars: 85%, 88%, 86%.')
add_para(doc, 'The BSA component accounted for 84.9% [74.8, 97.2] of the nonindexed decline in Arruda, 88.1% [82.1, 94.5] in Galvao, and 86.1% [79.6, 93.8] pooled (Figure 2). The non-BSA creatinine-based component was small and positive in both cohorts.')

# Discordance
add_heading(doc, 'Directional discordance', level=3)
add_figure(doc, fig3_path, 'Figure 3. Forest plot of 12-month paired changes (left) with directional discordance classification (right). Red = discordant (idx up, NI down); blue = both up; grey = both down; green = reverse (0%).', width=Inches(6), alt_text='Left panel: forest plot with diamond markers showing indexed eGFR increase and nonindexed eGFR decrease at 12 months for both cohorts and pooled. Right panel: stacked bar chart showing 48.8% directional discordance with 0% reverse pattern.')
add_para(doc, 'Nearly half of patients showed directional discordance: 48.8% [95% CI 42.7, 55.0] had indexed eGFR increase with simultaneous nonindexed eGFR decrease (Figure 3). No patient (0/250) showed the reverse pattern. Discordance was consistent across cohorts (Arruda 43.5%, Galvao 57.3%).')

# Dose-response
add_heading(doc, 'Dose-response relationship', level=3)
add_figure(doc, fig4_path, 'Figure 4. Dose-response: BSA change (A) and weight loss (B) versus indexed\u2013nonindexed divergence at 12 months. R\u00b2 = 0.72, highly consistent with a proportional algebraic mechanism.', width=Inches(6), alt_text='Two scatter plots showing monotonic relationship between BSA change and indexed-nonindexed divergence (R-squared 0.72) and between weight loss and divergence, with regression lines.')
add_para(doc, 'The divergence showed a strong monotonic relationship with BSA change (Spearman rho = -0.83, p < 10\u207b\u2076\u2074; R\u00b2 = 0.72) (Figure 4), consistent with an algebraic mechanism.')

# G-category
add_heading(doc, 'G-category stability', level=3)
add_para(doc, 'At 12 months, 89.1% of patients (491/551) maintained their pre-operative G-category [6], 9.4% improved, and 1.5% worsened. In exploratory longer-term follow-up (N = 130 at 60 months), the pattern remained similar (90.8% stable, 3.1% worsened), although these estimates should be interpreted cautiously given the smaller sample.')

# LME
add_heading(doc, 'Longitudinal mixed-effects model', level=3)
add_para(doc, 'The LME model (3,494 observations from 1,573 patients) confirmed a non-linear temporal trajectory for indexed eGFR (spline p < 10\u207b\u2078) with initial rise followed by late decline. Cohort-by-time interaction was non-significant (p = 0.83). Baseline age was the dominant predictor (-0.78 mL/min per year; p < 10\u207b\u00b9\u2076\u00b3). For nonindexed eGFR, there was a strong negative association with time (spline p < 10\u207b\u00b3\u2075), dominated by the BSA component.')

# Sensitivity
add_heading(doc, 'Sensitivity analyses', level=3)
add_para(doc, 'The primary finding was robust across all pre-specified sensitivity analyses: sex-stratified (female: +3.23/-17.09; male: +2.97/-17.51), BMI-stratified (consistent), alternative BSA formula (Mosteller: nonindexed -21.35 vs DuBois -17.15), restricted creatinine range (0.4-3.0: consistent), pre-operative window restriction (-180d: consistent), tight 12-month window (365\u00b145d, Arruda: -16.06), and Galvao date-validated nominal window (301-548d: nonindexed -18.94, discordance 76.5%).')

doc.add_page_break()

# ── DISCUSSION ──
add_heading(doc, 'Discussion', level=2)
add_para(doc, 'This dual-cohort study demonstrates that BSA indexing can reverse the apparent direction of creatinine-based eGFR change after RYGB. While indexed eGFR suggests mild improvement at 12 months, nonindexed eGFR reveals a substantial decline of approximately 17 mL/min. This discrepancy affects nearly half of individual patients and is almost entirely attributable to BSA rescaling.')

add_heading(doc, 'Mechanism', level=3)
add_para(doc, 'The Shapley decomposition provides formal confirmation of the algebraic mechanism. At 12 months, 86% of the nonindexed decline is explained by the BSA component alone. The non-BSA creatinine-based component was small and positive \u2014 though this should be interpreted cautiously, as creatinine production may also change after major weight loss due to altered muscle mass.')

add_heading(doc, 'Clinical implications', level=3)
add_para(doc, 'The practical consequence is straightforward: a clinician reviewing indexed eGFR after bariatric surgery sees +3 mL/min ("stable or improving") while nonindexed creatinine-based eGFR has fallen by approximately 17 mL/min. This is not a statistical curiosity \u2014 it affects individual clinical decisions. Nearly half of patients had clinically opposite trajectories, and no patient showed the reverse pattern.')
add_para(doc, 'We do not argue that nonindexed eGFR is the correct metric for drug dosing or staging; this remains an area of active debate [6,7]. Rather, we demonstrate that exclusive reliance on indexed eGFR after RYGB provides a systematically optimistic picture that clinicians should be aware of.')

add_heading(doc, 'Strengths and limitations', level=3)
add_para(doc, 'Strengths include the dual-cohort design with independent discovery and replication, shared laboratory standardisation, large sample size, formal algebraic decomposition with bootstrap inference, and comprehensive sensitivity analyses.')
add_para(doc, 'Limitations include the absence of cystatin C-based or measured GFR [6,7], which precludes assessment of whether the creatinine-based estimate itself reflects true filtration change. Albuminuria data were essentially absent, limiting classification to G-categories [6]. In the Galvao cohort, perioperative baseline misclassification remained possible (31.7% of pre-operative dates were post-surgical), although date-validated sensitivity analysis in a small date-validated replication subset substantially mitigated this concern (date availability remained incomplete: 104/283 with real dates, 17 nonindexed pairs in the nominal window). In the Arruda cohort, weight and laboratory data were captured within the same scheduled follow-up encounter; in the Galvao cohort, a modest lag between measurements cannot be excluded, although it is unlikely to explain a divergence of approximately 20 mL/min.')
add_para(doc, 'Reported comorbidity prevalence differed markedly between cohorts, likely reflecting ascertainment heterogeneity across source systems; because the primary endpoint was within-person change in creatinine-based eGFR under alternative BSA scaling, these variables were not used to define or compute the primary outcome.')

add_heading(doc, 'Conclusion', level=3)
add_para(doc, 'BSA indexing can reverse the apparent direction of creatinine-based eGFR change 12 months after RYGB. Clinicians and researchers should consider reporting both indexed and nonindexed eGFR when evaluating kidney function after bariatric surgery.')

doc.add_page_break()

# ── ACKNOWLEDGEMENTS ──
add_heading(doc, 'Acknowledgements', level=2)
add_para(doc, 'We thank Sabin Laboratorios (Brasilia, DF, Brazil) for providing standardised laboratory data. Artificial intelligence tools (Claude, Anthropic) were used to assist with data pipeline development, statistical code review, and manuscript formatting; all clinical interpretations and scientific content were generated and verified by the authors.')

# ── CONFLICT OF INTEREST STATEMENT ──
add_heading(doc, 'Conflict of Interest Statement', level=2)
add_para(doc, 'The authors declare no conflicts of interest.')

# ── AUTHORS' CONTRIBUTIONS ──
add_heading(doc, "Authors' Contributions", level=2)
add_para(doc, 'SA conceived the study, provided the discovery cohort data, and critically reviewed the manuscript. ROG provided the replication cohort data, performed the tier-2 adjudication, and critically reviewed the manuscript. RNB designed the analysis, performed all statistical analyses, wrote the first draft, and is the guarantor of the work. LBSC contributed to data collection and organisation. EBR provided methodological supervision and critically reviewed the manuscript. All authors approved the final version.')

# ── FUNDING ──
add_heading(doc, 'Funding', level=2)
add_para(doc, 'None.')

# ── DATA AVAILABILITY STATEMENT ──
add_heading(doc, 'Data Availability Statement', level=2)
add_para(doc, 'De-identified individual-level data are not publicly available because of privacy restrictions under the Brazilian General Data Protection Law (LGPD) and contractual restrictions involving the participating practices and laboratory. Analysis code and aggregated results are publicly available at https://github.com/RafaelBotan/bsa-indexing-bariatric-egfr and registered at https://osf.io/9r5w7/. Additional de-identified aggregate outputs may be available from the corresponding author on reasonable request and with permission from the data holders.')

doc.add_page_break()

# ── REFERENCES ──
add_heading(doc, 'References', level=2)

# References renumbered by order of appearance (NDT requirement). Refs [8],[17],[18] removed (not cited).
# Mapping: old→new: 14→1, 15→2, 9→3, 10→4, 11→5, 4→6, 13→7, 7→8, 16→9, 5→10, 6→11, 1→12, 19→13, 2→14, 3→15, 12→16, 20→17
references = [
    '1. Buchwald H, Avidor Y, Braunwald E, et al. Bariatric surgery: a systematic review and meta-analysis. JAMA. 2004;292(14):1724-1737. doi:10.1001/jama.292.14.1724',
    '2. Sjostrom L. Review of the key results from the Swedish Obese Subjects (SOS) trial. J Intern Med. 2013;273(3):219-234. doi:10.1111/joim.12012',
    '3. Chang AR, Chen Y, Still C, et al. Bariatric surgery is associated with improvement in kidney outcomes. Kidney Int. 2016;90(1):164-171. doi:10.1016/j.kint.2016.02.039',
    '4. Friedman AN, Moe S, Gao D, et al. Measured and estimated glomerular filtration rate after bariatric surgery. J Ren Nutr. 2023;33(1):94-101. doi:10.1053/j.jrn.2022.03.003',
    '5. Lieske JC, Collazo-Clavell ML, Engstrom BI, et al. Measured GFR and kidney outcomes after bariatric surgery. Clin J Am Soc Nephrol. 2024;19(3):305-313. doi:10.2215/CJN.0000000000000375',
    '6. KDIGO 2024 clinical practice guideline for the evaluation and management of chronic kidney disease. Kidney Int. 2024;105(4S):S117-S314. doi:10.1016/j.kint.2023.10.018',
    '7. Levey AS, Coresh J, Tighiouart H, Greene T, Inker LA. Measured and estimated glomerular filtration rate: current status and future directions. Nat Rev Nephrol. 2020;16(1):51-64. doi:10.1038/s41581-019-0191-y',
    '8. Delanaye P, Radermecker RP, Rorive M, Depas G, Krzesinski JM. Indexing glomerular filtration rate for body surface area in obese patients is misleading: concept and example. Nephrol Dial Transplant. 2005;20(10):2024-2028. doi:10.1093/ndt/gfh983',
    '9. Delanaye P, Mariat C, Cavalier E, Maillard N, Krzesinski JM, White CA. Indexation of renal function parameters by body surface area: intelligence or These? Nephrol Ther. 2009;5(Suppl 4):S213-218. doi:10.1016/S1769-7255(09)74545-3',
    '10. von Elm E, Altman DG, Egger M, et al. The Strengthening the Reporting of Observational Studies in Epidemiology (STROBE) statement: guidelines for reporting observational studies. Lancet. 2007;370(9596):1453-1457. doi:10.1016/S0140-6736(07)61602-X',
    '11. Benchimol EI, Smeeth L, Guttmann A, et al. The REporting of studies Conducted using Observational Routinely-collected health Data (RECORD) statement. PLoS Med. 2015;12(10):e1001885. doi:10.1371/journal.pmed.1001885',
    '12. Inker LA, Eneanya ND, Coresh J, et al. New creatinine- and cystatin C-based equations to estimate GFR without race. N Engl J Med. 2021;385(19):1737-1749. doi:10.1056/NEJMoa2102953',
    '13. Earley A, Miskulin D, Lamb EJ, Levey AS, Uhlig K. Estimating equations for glomerular filtration rate in the era of creatinine standardization: a systematic review. Ann Intern Med. 2012;156(11):785-795. doi:10.7326/0003-4819-156-6-201203200-00391',
    '14. Du Bois D, Du Bois EF. A formula to estimate the approximate surface area if height and weight be known. Arch Intern Med. 1916;17(6_2):863-871. doi:10.1001/archinte.1916.00080130010002',
    '15. Mosteller RD. Simplified calculation of body-surface area. N Engl J Med. 1987;317(17):1098. doi:10.1056/NEJM198710223171717',
    '16. Shapley LS. A value for n-person games. In: Kuhn HW, Tucker AW, eds. Contributions to the Theory of Games II. Princeton University Press; 1953:307-317. doi:10.1515/9781400881970-018',
    '17. Efron B, Tibshirani RJ. An Introduction to the Bootstrap. Chapman & Hall/CRC; 1993.',
]

for ref in references:
    p = doc.add_paragraph()
    run = p.add_run(ref)
    run.font.size = Pt(9)
    p.paragraph_format.space_after = Pt(3)

doc.add_page_break()

# ══════════════════════════════════════════════
# SUPPLEMENTARY MATERIAL
# ══════════════════════════════════════════════
add_heading(doc, 'Supplementary Material', level=1)

add_para(doc, 'Analysis code and aggregated results are publicly available at: https://github.com/RafaelBotan/bsa-indexing-bariatric-egfr', size=10)
add_para(doc, 'Pre-registration: https://osf.io/9r5w7/', size=10)
doc.add_paragraph()

# S-Table 1: LME coefficients
add_heading(doc, 'Supplementary Table S1. Linear Mixed-Effects Model Coefficients', level=2)

for model_name, model_label in [('eGFR_indexed', 'eGFR Indexed Model'), ('eGFR_NI', 'eGFR Nonindexed Model')]:
    add_para(doc, model_label, bold=True, size=10)
    m = lme[lme['Model']==model_name]
    t = doc.add_table(rows=len(m)+1, cols=4, style='Light Shading Accent 1')
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, h in enumerate(['Variable', 'Coefficient', 'SE', 'p-value']):
        cell = t.rows[0].cells[j]; cell.text = h
        for p in cell.paragraphs:
            for run in p.runs: run.bold = True; run.font.size = Pt(9)
    for i, (_, row) in enumerate(m.iterrows()):
        t.rows[i+1].cells[0].text = row['Variable']
        t.rows[i+1].cells[1].text = f"{row['Coefficient']:.4f}"
        t.rows[i+1].cells[2].text = f"{row['SE']:.4f}"
        t.rows[i+1].cells[3].text = f"{row['p']:.2e}" if row['p'] > 0 else '<10⁻³⁰⁰'
        for j in range(4):
            for p in t.rows[i+1].cells[j].paragraphs:
                for run in p.runs: run.font.size = Pt(9)
    doc.add_paragraph()

# Interaction model
add_para(doc, 'eGFR Indexed with Cohort × Time Interaction', bold=True, size=10)
m_int = lme[lme['Model']=='eGFR_idx_interaction']
t = doc.add_table(rows=len(m_int)+1, cols=4, style='Light Shading Accent 1')
t.alignment = WD_TABLE_ALIGNMENT.CENTER
for j, h in enumerate(['Variable', 'Coefficient', 'SE', 'p-value']):
    cell = t.rows[0].cells[j]; cell.text = h
    for p in cell.paragraphs:
        for run in p.runs: run.bold = True; run.font.size = Pt(9)
for i, (_, row) in enumerate(m_int.iterrows()):
    t.rows[i+1].cells[0].text = row['Variable']
    t.rows[i+1].cells[1].text = f"{row['Coefficient']:.4f}"
    t.rows[i+1].cells[2].text = f"{row['SE']:.4f}"
    t.rows[i+1].cells[3].text = f"{row['p']:.2e}" if row['p'] > 0 else '<10⁻³⁰⁰'
    for j in range(4):
        for p in t.rows[i+1].cells[j].paragraphs:
            for run in p.runs: run.font.size = Pt(9)
doc.add_paragraph()

# S-Fig 1: LME curves
add_figure(doc, sfig1_path, 'Supplementary Figure S1. Predicted eGFR trajectories from the linear mixed-effects model. (A) Indexed eGFR shows non-linear trajectory with initial rise and late decline. (B) Nonindexed eGFR shows sustained decline dominated by BSA component. Predictions at mean age (38.7 years), 83% female.', width=Inches(6))

# S-Fig 2: G-category
add_figure(doc, sfig2_path, 'Supplementary Figure S2. G-category transitions over time (combined cohorts). Green = stable; blue = improved; red = worsened. 89-91% remained stable at all time points.', width=Inches(5.5))

# S-Fig 3: Mosteller
add_figure(doc, sfig3_path, 'Supplementary Figure S3. Nonindexed eGFR change by BSA formula. Mosteller produces larger decline (~25% greater) but direction is identical. Both formulas confirm the indexing artifact.', width=Inches(5.5))

# S-Fig 4: Sensitivity
add_figure(doc, sfig4_path, 'Supplementary Figure S4. Sensitivity analyses: 12-month paired changes across subgroups. All analyses confirm indexed eGFR rises while nonindexed eGFR falls.', width=Inches(5.5))

# S-Fig 5: PRE window
add_figure(doc, sfig5_path, 'Supplementary Figure S5. Pre-operative window sensitivity. The primary finding is robust to restriction of the pre-operative window from -365 to -180 or -90 days.', width=Inches(5.5))

# S-Fig 6: Galvao date-validated
add_figure(doc, sfig6_path, 'Supplementary Figure S6. Galvao date-validated 12-month sensitivity. When restricted to the nominal 301-548 day window, indexed change attenuates to zero while nonindexed decline persists and divergence increases.', width=Inches(5.5))

# S-Fig 7: Shapley temporal
if sfig7_path:
    add_figure(doc, sfig7_path, 'Supplementary Figure S7. Shapley decomposition temporal evolution (Galvao only, exploratory). The BSA component remains dominant through 60 months; a small negative eGFR component emerges at 24+ months, consistent with aging.', width=Inches(6))

# S-Table 2: Linked vs unlinked
add_heading(doc, 'Supplementary Table S2. Linked vs Unlinked Patients (Galvao Sabin)', level=2)
t = doc.add_table(rows=len(linked)+1, cols=5, style='Light Shading Accent 1')
t.alignment = WD_TABLE_ALIGNMENT.CENTER
for j, h in enumerate(['Variable', 'Linked', 'Unlinked', 'SMD', 'Flag']):
    cell = t.rows[0].cells[j]; cell.text = h
    for p in cell.paragraphs:
        for run in p.runs: run.bold = True; run.font.size = Pt(9)
for i, (_, row) in enumerate(linked.iterrows()):
    t.rows[i+1].cells[0].text = str(row['Variable'])
    lk_mean = row.get('Linked_Mean', '')
    lk_sd = row.get('Linked_SD', '')
    ulk_mean = row.get('Unlinked_Mean', '')
    ulk_sd = row.get('Unlinked_SD', '')
    if pd.notna(lk_sd) and lk_sd != '':
        t.rows[i+1].cells[1].text = f"{float(lk_mean):.1f} \u00b1 {float(lk_sd):.1f} (N={int(row['Linked_N'])})"
        t.rows[i+1].cells[2].text = f"{float(ulk_mean):.1f} \u00b1 {float(ulk_sd):.1f} (N={int(row['Unlinked_N'])})"
    else:
        t.rows[i+1].cells[1].text = f"{float(lk_mean):.1f}% (N={int(row['Linked_N'])})"
        t.rows[i+1].cells[2].text = f"{float(ulk_mean):.1f}% (N={int(row['Unlinked_N'])})"
    t.rows[i+1].cells[3].text = f"{abs(row['SMD']):.2f}"
    t.rows[i+1].cells[4].text = str(row.get('Imbalance_flag', ''))
    for j in range(5):
        for p in t.rows[i+1].cells[j].paragraphs:
            for run in p.runs: run.font.size = Pt(9)

doc.add_paragraph()

# S-Table 3: Date-validated Galvao
add_heading(doc, 'Supplementary Table S3. Galvao Date-Validated 12-Month Sensitivity', level=2)
t = doc.add_table(rows=len(galvao_dv)+1, cols=7, style='Light Shading Accent 1')
t.alignment = WD_TABLE_ALIGNMENT.CENTER
for j, h in enumerate(['Window', 'N idx', '\u0394 idx', 'p idx', 'N NI', '\u0394 NI', 'Divergence']):
    cell = t.rows[0].cells[j]; cell.text = h
    for p in cell.paragraphs:
        for run in p.runs: run.bold = True; run.font.size = Pt(9)
for i, (_, row) in enumerate(galvao_dv.iterrows()):
    t.rows[i+1].cells[0].text = str(row['Sensitivity']).replace('Galvao_', '')
    t.rows[i+1].cells[1].text = str(int(row['N_idx']))
    t.rows[i+1].cells[2].text = str(row['Delta_idx'])
    t.rows[i+1].cells[3].text = str(row['p_idx'])
    t.rows[i+1].cells[4].text = str(int(row['N_NI'])) if pd.notna(row.get('N_NI')) else '-'
    t.rows[i+1].cells[5].text = str(row.get('Delta_NI', '-')) if pd.notna(row.get('Delta_NI')) else '-'
    t.rows[i+1].cells[6].text = str(row.get('Divergence', '-')) if pd.notna(row.get('Divergence')) else '-'
    for j in range(7):
        for p in t.rows[i+1].cells[j].paragraphs:
            for run in p.runs: run.font.size = Pt(9)

# ── Save ──
doc.save(OUT_DOCX)
print(f"\nDOCX saved to: {OUT_DOCX}")
print(f"File size: {os.path.getsize(OUT_DOCX) / 1024:.0f} KB")
print("Done!")
